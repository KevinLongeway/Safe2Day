"""
Create sample S2D OHA - Admin document with Go_Auto1 logo
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Create document
doc = Document()

# Add logo to header
header = doc.sections[0].header
header_para = header.paragraphs[0]
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

logo_path = Path('Client_Logos/Go_Auto1.png')
if logo_path.exists():
    run = header_para.add_run()
    run.add_picture(str(logo_path), width=Inches(2.0))
else:
    header_para.add_run('Go_Auto1 Logo')

# Add title
doc.add_heading('Occupational Health Assessment - Admin', 0)

# Add form fields
doc.add_paragraph('Employee Name: ________________________')
doc.add_paragraph('Date: ________________________')
doc.add_paragraph('Position: ________________________')
doc.add_paragraph()

# Assessment section
doc.add_heading('Assessment Details', 1)
doc.add_paragraph('This is a sample S2D Occupational Health Assessment document for administrative staff.')
doc.add_paragraph()

# Health screening section
doc.add_heading('Health Screening', 1)
doc.add_paragraph('☐ Vision Test')
doc.add_paragraph('☐ Hearing Test')
doc.add_paragraph('☐ Blood Pressure')
doc.add_paragraph('☐ General Physical')
doc.add_paragraph()

# Signature section
doc.add_heading('Signature', 1)
doc.add_paragraph('Healthcare Provider: ________________________')
doc.add_paragraph('Date: ________________________')

# Add logo to footer
footer = doc.sections[0].footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

if logo_path.exists():
    run = footer_para.add_run()
    run.add_picture(str(logo_path), width=Inches(1.5))
else:
    footer_para.add_run('Go_Auto1 Logo')

# Save document
output_path = Path('Forms/S2D OHA - Admin.docx')
doc.save(str(output_path))

print(f'✓ Created: {output_path}')
print(f'  - Header logo: {logo_path.name}')
print(f'  - Footer logo: {logo_path.name}')
