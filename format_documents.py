"""
format_documents.py - Create client documents with custom logos
NEVER modifies source documents in Forms folder.
Creates copies in Client_Forms folder with selected logo, then generates PDFs.
"""

import os
import shutil
from pathlib import Path
import json
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from PIL import Image


def load_selected_logo():
    """
    Load the selected logo from the selection file.
    
    Returns:
        Path to selected logo or None
    """
    selection_file = Path(__file__).parent / "selected_logo.json"
    
    if not selection_file.exists():
        print("Error: No logo selected. Please run select_logo.py first.")
        return None
    
    with open(selection_file, 'r') as f:
        selection_data = json.load(f)
    
    logo_path = Path(selection_data['logo_path'])
    
    if not logo_path.exists():
        print(f"Error: Selected logo not found at {logo_path}")
        return None
    
    return logo_path


def load_logo_positions():
    """
    Load logo position data from JSON file.
    
    Returns:
        Dictionary of logo positions or None
    """
    data_file = Path(__file__).parent / "logo_positions.json"
    
    if not data_file.exists():
        print("Error: No logo position data found. Please run markup_document.py first.")
        return None
    
    with open(data_file, 'r') as f:
        position_data = json.load(f)
    
    return position_data


def replace_logo_in_header(doc, logo_path, logo_positions):
    """
    Replace existing logo in document header with new logo.
    
    Args:
        doc: Document object
        logo_path: Path to the new logo image
        logo_positions: List of logo position dictionaries
    
    Returns:
        Number of logos replaced
    """
    if not logo_positions:
        return 0
    
    replaced_count = 0
    
    for pos in logo_positions:
        try:
            if pos['type'] == 'header':
                section_idx = pos.get('section_index', 0)
                para_idx = pos['para_index']
                
                if section_idx < len(doc.sections):
                    section = doc.sections[section_idx]
                    if section.header and para_idx < len(section.header.paragraphs):
                        paragraph = section.header.paragraphs[para_idx]
                        
                        # Remove all existing images from runs
                        for run in paragraph.runs:
                            if hasattr(run._element, 'drawing_lst'):
                                # Remove all drawings (images)
                                for drawing in list(run._element.drawing_lst):
                                    run._element.remove(drawing)
                        
                        # Clear all runs and add new logo
                        paragraph.clear()
                        run = paragraph.add_run()
                        run.add_picture(str(logo_path), width=Inches(2.0))
                        
                        # Restore alignment
                        if pos.get('alignment'):
                            paragraph.alignment = pos['alignment']
                        
                        replaced_count += 1
                        print(f"  - Replaced logo in header section {section_idx}")
            
            elif pos['type'] == 'footer':
                section_idx = pos.get('section_index', 0)
                para_idx = pos['para_index']
                
                if section_idx < len(doc.sections):
                    section = doc.sections[section_idx]
                    if section.footer and para_idx < len(section.footer.paragraphs):
                        paragraph = section.footer.paragraphs[para_idx]
                        
                        # Remove all existing images from runs
                        for run in paragraph.runs:
                            if hasattr(run._element, 'drawing_lst'):
                                for drawing in list(run._element.drawing_lst):
                                    run._element.remove(drawing)
                        
                        # Clear all runs and add new logo
                        paragraph.clear()
                        run = paragraph.add_run()
                        run.add_picture(str(logo_path), width=Inches(1.5))
                        
                        # Restore alignment
                        if pos.get('alignment'):
                            paragraph.alignment = pos['alignment']
                        
                        replaced_count += 1
                        print(f"  - Replaced logo in footer section {section_idx}")
        
        except Exception as e:
            print(f"    Warning: Could not replace logo at position {pos}: {e}")
    
    return replaced_count


def convert_to_pdf(doc_path, output_folder):
    """
    Convert a Word document to PDF.
    
    Args:
        doc_path: Path to the Word document
        output_folder: Path to output folder for PDF
        
    Returns:
        Path to created PDF or None
    """
    try:
        pdf_path = output_folder / (doc_path.stem + ".pdf")
        convert(str(doc_path), str(pdf_path))
        print(f"  ✓ PDF created: {pdf_path.name}")
        return pdf_path
    except Exception as e:
        print(f"  ✗ Error converting to PDF: {e}")
        return None


def process_documents():
    """
    Main function: Copy source documents to Client_Forms, replace logos, and create PDFs.
    NEVER modifies source documents in Forms folder.
    """
    # Load selected logo
    logo_path = load_selected_logo()
    if not logo_path:
        return
    
    print(f"Using logo: {logo_path.name}\n")
    
    # Load logo position data
    position_data = load_logo_positions()
    if not position_data:
        return
    
    # Setup folders
    forms_folder = Path(__file__).parent / "Forms"
    client_forms_folder = Path(__file__).parent / "Client_Forms"
    pdf_folder = Path(__file__).parent / "PDFs"
    
    # Create output folders
    client_forms_folder.mkdir(exist_ok=True)
    pdf_folder.mkdir(exist_ok=True)
    
    if not forms_folder.exists():
        print(f"Error: Forms folder not found at {forms_folder}")
        return
    
    # Find all S2D Word documents in source folder
    word_files = list(forms_folder.glob("S2D*.docx"))
    
    if not word_files:
        print("No Word documents found with 'S2D' prefix in Forms folder.")
        return
    
    print(f"Processing {len(word_files)} document(s)...\n")
    
    success_count = 0
    
    for source_doc_path in word_files:
        doc_name = source_doc_path.name
        print(f"Processing: {doc_name}")
        
        try:
            # Step 1: Copy source document to Client_Forms (NEVER modify source!)
            client_doc_path = client_forms_folder / doc_name
            shutil.copy2(source_doc_path, client_doc_path)
            print(f"  ✓ Copied to Client_Forms")
            
            # Step 2: Open the COPY and replace logo
            doc = Document(client_doc_path)
            positions = position_data.get(doc_name, [])
            
            if positions:
                replaced = replace_logo_in_header(doc, logo_path, positions)
                doc.save(client_doc_path)
                print(f"  ✓ Replaced {replaced} logo(s)")
            else:
                print(f"  - No logo positions found (skipping logo replacement)")
            
            # Step 3: Convert the client copy to PDF
            pdf_path = convert_to_pdf(client_doc_path, pdf_folder)
            
            if pdf_path:
                success_count += 1
            
            print()
        
        except Exception as e:
            print(f"  ✗ Error processing {doc_name}: {e}\n")
    
    print("=" * 60)
    print(f"✓ Successfully processed {success_count}/{len(word_files)} document(s)")
    print(f"✓ Word documents saved to: {client_forms_folder}")
    print(f"✓ PDFs saved to: {pdf_folder}")
    print(f"\n⚠ NOTE: Source documents in Forms folder were NOT modified")


if __name__ == "__main__":
    print("=" * 60)
    print("Safe2Day Document Formatter")
    print("=" * 60)
    print()
    process_documents()
