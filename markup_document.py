"""
markup_document.py - Analyze S2D logos in Word documents
This script ANALYZES all Word documents in the Forms folder with prefix "S2D",
identifies logo positions in headers for future logo replacement.
NOTE: This script NEVER modifies source documents in Forms folder.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json


def analyze_logo_positions(doc_path, placeholder_data):
    """
    Analyze a Word document to find logo positions in headers.
    Does NOT modify the document.
    
    Args:
        doc_path: Path to the Word document
        placeholder_data: Dictionary to store placeholder positions
    """
    doc = Document(doc_path)
    doc_name = os.path.basename(doc_path)
    placeholder_data[doc_name] = []
    
    print(f"Analyzing: {doc_name}")
    
    # Check headers for images (logos)
    for section_idx, section in enumerate(doc.sections):
        # Check header
        if section.header:
            for para_idx, paragraph in enumerate(section.header.paragraphs):
                for run_idx, run in enumerate(paragraph.runs):
                    if hasattr(run._element, 'drawing_lst') and len(run._element.drawing_lst) > 0:
                        # Found a logo in header
                        placeholder_data[doc_name].append({
                            'type': 'header',
                            'section_index': section_idx,
                            'para_index': para_idx,
                            'run_index': run_idx,
                            'alignment': paragraph.alignment
                        })
                        print(f"  - Found logo in header section {section_idx}, paragraph {para_idx}")
        
        # Check footer
        if section.footer:
            for para_idx, paragraph in enumerate(section.footer.paragraphs):
                for run_idx, run in enumerate(paragraph.runs):
                    if hasattr(run._element, 'drawing_lst') and len(run._element.drawing_lst) > 0:
                        # Found a logo in footer
                        placeholder_data[doc_name].append({
                            'type': 'footer',
                            'section_index': section_idx,
                            'para_index': para_idx,
                            'run_index': run_idx,
                            'alignment': paragraph.alignment
                        })
                        print(f"  - Found logo in footer section {section_idx}, paragraph {para_idx}")
    
    if not placeholder_data[doc_name]:
        print(f"  - No logos found in headers/footers")
    
    print()


def analyze_forms_folder():
    """
    Analyze all Word documents in Forms folder that start with 'S2D'.
    Does NOT modify any source documents.
    """
    forms_folder = Path(__file__).parent / "Forms"
    
    if not forms_folder.exists():
        print(f"Error: Forms folder not found at {forms_folder}")
        return
    
    # Dictionary to store all logo positions
    placeholder_data = {}
    
    # Find all Word documents with S2D prefix
    word_files = list(forms_folder.glob("S2D*.docx"))
    
    if not word_files:
        print("No Word documents found with 'S2D' prefix in Forms folder.")
        return
    
    print(f"Found {len(word_files)} document(s) to analyze.\n")
    
    for doc_path in word_files:
        try:
            analyze_logo_positions(str(doc_path), placeholder_data)
        except Exception as e:
            print(f"Error analyzing {doc_path.name}: {e}\n")
    
    # Save logo position data to JSON file
    data_file = Path(__file__).parent / "logo_positions.json"
    with open(data_file, 'w') as f:
        json.dump(placeholder_data, f, indent=2)
    
    print(f"✓ Logo position data saved to: {data_file}")
    print(f"✓ Analyzed {len(placeholder_data)} document(s)")
    print(f"\n⚠ NOTE: Source documents in Forms folder were NOT modified")


if __name__ == "__main__":
    print("=" * 60)
    print("Safe2Day Logo Analysis Tool")
    print("=" * 60)
    print()
    analyze_forms_folder()
